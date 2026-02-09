#!/usr/bin/env python3
"""
Generate professionally formatted DOCX and PDF from draft_report.md
Uses python-docx for fine-grained control over fonts, spacing, and layout.
"""

import re
import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

SCRIPT_DIR = Path(__file__).parent
MD_FILE = SCRIPT_DIR / "draft_report.md"
IMAGES_DIR = SCRIPT_DIR / "images"
OUTPUT_DOCX = SCRIPT_DIR / "NextTrack_Draft_Report.docx"
OUTPUT_PDF = SCRIPT_DIR / "NextTrack_Draft_Report.pdf"


def setup_styles(doc):
    """Configure document styles for professional academic formatting."""
    # --- Page margins ---
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # --- Normal style ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.space_before = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    # --- Title ---
    title_style = doc.styles['Title']
    title_style.font.name = 'Helvetica Neue'
    title_style.font.size = Pt(26)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    title_style.paragraph_format.space_after = Pt(4)
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Subtitle ---
    subtitle_style = doc.styles['Subtitle']
    subtitle_style.font.name = 'Helvetica Neue'
    subtitle_style.font.size = Pt(14)
    subtitle_style.font.italic = True
    subtitle_style.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)
    subtitle_style.paragraph_format.space_after = Pt(12)
    subtitle_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Heading 1 (Chapter headings) ---
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Helvetica Neue'
    h1.font.size = Pt(20)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.page_break_before = True

    # --- Heading 2 ---
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Helvetica Neue'
    h2.font.size = Pt(15)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0x2C, 0x5F, 0x8A)
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(8)

    # --- Heading 3 ---
    h3 = doc.styles['Heading 3']
    h3.font.name = 'Helvetica Neue'
    h3.font.size = Pt(13)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0x3A, 0x7C, 0xB8)
    h3.paragraph_format.space_before = Pt(14)
    h3.paragraph_format.space_after = Pt(6)

    # --- Heading 4 ---
    h4 = doc.styles['Heading 4']
    h4.font.name = 'Helvetica Neue'
    h4.font.size = Pt(11.5)
    h4.font.bold = True
    h4.font.italic = True
    h4.font.color.rgb = RGBColor(0x3A, 0x7C, 0xB8)
    h4.paragraph_format.space_before = Pt(10)
    h4.paragraph_format.space_after = Pt(4)

    # --- Code block style ---
    if 'Code Block' not in [s.name for s in doc.styles]:
        code_style = doc.styles.add_style('Code Block', WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = 'Courier New'
        code_style.font.size = Pt(8.5)
        code_style.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
        code_style.paragraph_format.space_before = Pt(6)
        code_style.paragraph_format.space_after = Pt(6)
        code_style.paragraph_format.line_spacing = Pt(12)
        code_style.paragraph_format.left_indent = Cm(0.5)

    # --- Figure caption style ---
    if 'Caption' in [s.name for s in doc.styles]:
        cap = doc.styles['Caption']
        cap.font.name = 'Times New Roman'
        cap.font.size = Pt(10)
        cap.font.italic = True
        cap.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)
        cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_before = Pt(4)
        cap.paragraph_format.space_after = Pt(12)

    return doc


def add_horizontal_rule(doc):
    """Add a subtle horizontal rule."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('─' * 60)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    run.font.size = Pt(8)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)


def add_code_block(doc, code_lines):
    """Add a formatted code block with background shading."""
    for line in code_lines:
        p = doc.add_paragraph(line, style='Code Block')
        # Add light grey background shading
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>')
        p.paragraph_format.element.get_or_add_pPr().append(shading)


def add_table_from_md(doc, header_line, separator_line, data_lines):
    """Parse markdown table and create a formatted Word table."""
    def parse_row(line):
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        return cells

    headers = parse_row(header_line)
    rows = [parse_row(line) for line in data_lines if line.strip()]

    num_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Light Grid Accent 1'

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.name = 'Helvetica Neue'
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Dark blue header background
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1B3A5C" w:val="clear"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            if col_idx < num_cols:
                cell = table.rows[row_idx + 1].cells[col_idx]
                cell.text = ''
                p = cell.paragraphs[0]
                # Handle bold text in cells
                parts = re.split(r'\*\*(.*?)\*\*', cell_text)
                for j, part in enumerate(parts):
                    if j % 2 == 1:
                        run = p.add_run(part)
                        run.bold = True
                    else:
                        run = p.add_run(part)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                # Alternating row colors
                if row_idx % 2 == 1:
                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="EDF2F7" w:val="clear"/>')
                    cell._tc.get_or_add_tcPr().append(shading)

    # Add spacing after table
    doc.add_paragraph()
    return table


def add_image(doc, image_path, caption_text=None, width=Inches(5.5)):
    """Add an image with optional caption."""
    if os.path.exists(image_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(image_path), width=width)

        if caption_text:
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cap.add_run(caption_text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)
            cap.paragraph_format.space_after = Pt(12)
    else:
        p = doc.add_paragraph(f"[Image not found: {image_path}]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def process_inline_formatting(paragraph, text):
    """Apply inline markdown formatting (bold, italic, code, links) to a paragraph."""
    # Pattern to match bold, italic, inline code, and links
    pattern = r'(\*\*\*(.*?)\*\*\*|\*\*(.*?)\*\*|\*(.*?)\*|`(.*?)`|\[(.*?)\]\((.*?)\))'

    last_end = 0
    for match in re.finditer(pattern, text):
        # Add text before this match
        if match.start() > last_end:
            run = paragraph.add_run(text[last_end:match.start()])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

        if match.group(2):  # Bold italic ***text***
            run = paragraph.add_run(match.group(2))
            run.bold = True
            run.italic = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
        elif match.group(3):  # Bold **text**
            run = paragraph.add_run(match.group(3))
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
        elif match.group(4):  # Italic *text*
            run = paragraph.add_run(match.group(4))
            run.italic = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
        elif match.group(5):  # Inline code `text`
            run = paragraph.add_run(match.group(5))
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        elif match.group(6):  # Link [text](url)
            run = paragraph.add_run(match.group(6))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x2C, 0x5F, 0x8A)

        last_end = match.end()

    # Add remaining text
    if last_end < len(text):
        run = paragraph.add_run(text[last_end:])
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)


def parse_and_convert(md_text, doc):
    """Parse markdown and convert to formatted DOCX elements."""
    lines = md_text.split('\n')
    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_header = ""
    table_separator = ""
    table_data = []
    first_h1 = True

    # Skip to first meaningful content (skip front matter)
    while i < len(lines):
        line = lines[i]

        # --- Code blocks ---
        if line.strip().startswith('```'):
            if in_code_block:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code_block = False
                i += 1
                continue
            else:
                in_code_block = True
                i += 1
                continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # --- Tables ---
        if '|' in line and not line.strip().startswith('```'):
            # Check if next line is a separator
            if i + 1 < len(lines) and re.match(r'^[\s|:-]+$', lines[i + 1]):
                if not in_table:
                    in_table = True
                    table_header = line
                    table_separator = lines[i + 1]
                    table_data = []
                    i += 2
                    continue
            elif in_table:
                table_data.append(line)
                i += 1
                continue

        if in_table and ('|' not in line or line.strip() == ''):
            # End of table
            add_table_from_md(doc, table_header, table_separator, table_data)
            in_table = False
            table_data = []
            if line.strip() == '':
                i += 1
                continue

        # --- Horizontal rules ---
        if line.strip() in ['---', '***', '___']:
            i += 1
            continue

        # --- Empty lines ---
        if line.strip() == '':
            i += 1
            continue

        # --- Images ---
        img_match = re.match(r'!\[(.*?)\]\((.*?)\)', line.strip())
        if img_match:
            alt_text = img_match.group(1)
            img_path = img_match.group(2)
            full_path = SCRIPT_DIR / img_path
            add_image(doc, full_path, alt_text, width=Inches(5.0))
            i += 1
            continue

        # --- Figure captions (italic lines starting with *Figure) ---
        if line.strip().startswith('*Figure') and line.strip().endswith('*'):
            caption_text = line.strip().strip('*')
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(caption_text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)
            p.paragraph_format.space_after = Pt(12)
            i += 1
            continue

        # --- Headings ---
        h_match = re.match(r'^(#{1,4})\s+(.*)', line)
        if h_match:
            level = len(h_match.group(1))
            heading_text = h_match.group(2)

            if level == 1:
                # First H1 is the title
                if first_h1:
                    doc.add_paragraph(heading_text, style='Title')
                    first_h1 = False
                else:
                    p = doc.add_paragraph(heading_text, style='Heading 1')
            elif level == 2:
                # Check if it's the subtitle
                if heading_text.startswith('CM3035'):
                    doc.add_paragraph(heading_text, style='Subtitle')
                else:
                    doc.add_paragraph(heading_text, style='Heading 2')
            elif level == 3:
                doc.add_paragraph(heading_text, style='Heading 3')
            elif level == 4:
                doc.add_paragraph(heading_text, style='Heading 4')
            i += 1
            continue

        # --- Ordered lists ---
        ol_match = re.match(r'^(\d+)\.\s+(.*)', line)
        if ol_match:
            text = ol_match.group(2)
            p = doc.add_paragraph(style='List Number')
            process_inline_formatting(p, text)
            p.paragraph_format.left_indent = Cm(1.27)
            p.paragraph_format.space_after = Pt(3)
            i += 1
            continue

        # --- Unordered lists ---
        ul_match = re.match(r'^[-*+]\s+(.*)', line)
        if ul_match:
            text = ul_match.group(1)
            p = doc.add_paragraph(style='List Bullet')
            process_inline_formatting(p, text)
            p.paragraph_format.left_indent = Cm(1.27)
            p.paragraph_format.space_after = Pt(3)
            i += 1
            continue

        # --- Bold paragraph headers (like **Privacy Erosion**) ---
        if line.strip().startswith('**') and line.strip().endswith('**'):
            text = line.strip().strip('*')
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            p.paragraph_format.space_before = Pt(8)
            i += 1
            continue

        # --- Regular paragraphs ---
        # Collect continuation lines for the paragraph
        para_text = line.strip()
        p = doc.add_paragraph()
        process_inline_formatting(p, para_text)
        i += 1

    # Handle any remaining table
    if in_table and table_data:
        add_table_from_md(doc, table_header, table_separator, table_data)


def add_cover_page(doc):
    """Add a professional cover page."""
    # Add some spacing
    for _ in range(4):
        doc.add_paragraph()

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('NextTrack')
    run.font.name = 'Helvetica Neue'
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('A Privacy-Focused Music Recommendation API')
    run.font.name = 'Helvetica Neue'
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)
    p.paragraph_format.space_after = Pt(36)

    # Horizontal line
    add_horizontal_rule(doc)

    # Course info
    info_items = [
        ('Course', 'CM3035 Advanced Web Design'),
        ('Project Type', 'RESTful API Development'),
        ('Date', 'February 2026'),
        ('Word Count', '~8,700 words (excl. references, figures & tables)'),
    ]

    for label, value in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'{label}: ')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x6A, 0x6A, 0x6A)
        run = p.add_run(value)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        p.paragraph_format.space_after = Pt(2)

    # Add page break
    doc.add_page_break()


def add_toc_placeholder(doc):
    """Add a Table of Contents placeholder with field code."""
    p = doc.add_paragraph('Table of Contents', style='Heading 1')
    p.style.paragraph_format.page_break_before = False  # Don't page break for TOC

    # Add instruction to update TOC
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('[Right-click and select "Update Field" to generate Table of Contents]')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Add TOC field
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fldChar1)

    run2 = paragraph.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>')
    run2._r.append(instrText)

    run3 = paragraph.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run3._r.append(fldChar2)

    run4 = paragraph.add_run('Press F9 or right-click to update this table of contents')
    run4.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    run5 = paragraph.add_run()
    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run5._r.append(fldChar3)

    doc.add_page_break()


def add_page_numbers(doc):
    """Add page numbers to footer."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Page number field
        run = p.add_run()
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._r.append(fldChar1)

        run2 = p.add_run()
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run2._r.append(instrText)

        run3 = p.add_run()
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
        run3._r.append(fldChar2)

        run4 = p.add_run('1')
        run4.font.name = 'Times New Roman'
        run4.font.size = Pt(9)
        run4.font.color.rgb = RGBColor(0x6A, 0x6A, 0x6A)

        run5 = p.add_run()
        fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run5._r.append(fldChar3)


def main():
    """Main function to generate DOCX from markdown."""
    print("📄 Reading draft report...")
    md_text = MD_FILE.read_text(encoding='utf-8')

    print("📝 Creating Word document...")
    doc = Document()
    setup_styles(doc)

    # Cover page
    print("   → Adding cover page...")
    add_cover_page(doc)

    # Table of contents
    print("   → Adding table of contents...")
    add_toc_placeholder(doc)

    # Main content (skip the first title and metadata as cover page handles it)
    # Remove the metadata block at the top
    content_start = md_text.find('## Table of Contents')
    if content_start != -1:
        # Skip the markdown TOC and start from Chapter 1
        chapter_start = md_text.find('# Chapter 1')
        if chapter_start != -1:
            md_content = md_text[chapter_start:]
        else:
            md_content = md_text[content_start:]
    else:
        md_content = md_text

    print("   → Converting content...")
    parse_and_convert(md_content, doc)

    # Page numbers
    print("   → Adding page numbers...")
    add_page_numbers(doc)

    # Save DOCX
    print(f"💾 Saving DOCX to: {OUTPUT_DOCX}")
    doc.save(str(OUTPUT_DOCX))
    print(f"✅ DOCX generated successfully: {OUTPUT_DOCX.name}")
    print(f"   File size: {OUTPUT_DOCX.stat().st_size / 1024:.1f} KB")

    return str(OUTPUT_DOCX)


if __name__ == '__main__':
    docx_path = main()
